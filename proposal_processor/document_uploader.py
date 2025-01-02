import os
import io
from typing import List
from langchain.schema import Document
from langchain_google_vertexai import VertexAIEmbeddings
from langchain_community.vectorstores import Chroma
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

def upload_documents(directory_path: str) -> None:
    """
    Upload documents from a directory to Chroma vector store.
    
    Args:
        directory_path (str): Path to directory containing documents
    """
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if os.path.isfile(file_path):
            try:
                # Open files in binary mode
                with open(file_path, 'rb') as f:
                    binary_content = f.read()
                    
                    # Handle different file types
                    if filename.lower().endswith('.pdf'):
                        reader = PdfReader(io.BytesIO(binary_content))
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text()
                    
                    elif filename.lower().endswith(('.doc', '.docx')):
                        doc = DocxDocument(io.BytesIO(binary_content))
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    
                    else:  # Try different encodings for text files
                        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
                        text = None
                        for encoding in encodings:
                            try:
                                text = binary_content.decode(encoding)
                                break
                            except UnicodeDecodeError:
                                continue
                        if text is None:
                            print(f"Warning: Could not decode {filename} with any supported encoding")
                            continue
                    
                    documents.append(Document(page_content=text, metadata={"source": filename}))
            except Exception as e:
                print(f"Error processing {filename}: {str(e)}")
                continue

    # Initialize embeddings and vector store
    embeddings = VertexAIEmbeddings(
        model_name="textembedding-gecko"
    )
    
    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory="chroma_db"
    )
    
    vectordb.persist()